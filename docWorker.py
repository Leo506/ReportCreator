from curses import doupdate
from http import server
from docx.shared import Inches
from docx.shared import Pt
import userInput
import os

def addHead(document, settings):
    head = userInput.getString(settings["helpText"])

    h = document.add_heading(head, 0)

    h.runs[0].font.size = Pt(int(settings["fontSize"]))


def addParagraph(document, settings):
    paragraph = userInput.getString(settings["helpText"])

    p = document.add_paragraph(paragraph)

    p.runs[0].font.size = Pt(int(settings["fontSize"]))


def addFileText(document, settings):
    try:

        url = userInput.getString(settings["helpText"])
        while not os.path.exists(url):
            print("Can't find file")
            url = userInput.getString(settings["helpText"])    

        with open(url, 'r', encoding="utf-8") as file:
            text = file.read()
            t = document.add_paragraph(text)
            t.runs[0].font.size = Pt(int(settings["fontSize"]))
            
    except Exception as e:
        addFileText(document, settings)


def addPicture(document, settings):
    url = userInput.getString(settings["helpText"])
    while not os.path.exists(url):
        print("Can't find file")
        url = userInput.getString(settings["helpText"])

    document.add_picture(url, width=Inches(3))